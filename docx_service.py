import datetime
import os
from copy import deepcopy

from docx import *
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import _Cell
from consts import *
from data_classes import *


class DocxService:

    def __init__(self, docx_path: str, multiplicity_times_dict: dict):
        self.docx_path = docx_path
        self.document = None
        self.patient: Patient = None
        self.multiplicity_times_dict = multiplicity_times_dict
        self.max_date_range: list[datetime.datetime] = []
        
    def generate_document(self, patient: Patient):
        self.patient = patient
        self.document = Document(self.docx_path)
        self.set_fio_and_ward()
        self.set_mode_and_diet()
        self.set_date_range()
        self.set_prescriptions()
        self.set_surveys()
        self.add_nurse_tables()
        self.save_document()

    def set_fio_and_ward(self):
        cell = self.document.tables[0].cell(2, 0)  # ЛИСТОК ЛІКАРСЬКИХ ПРИЗНАЧЕНЬ
        for par in cell.paragraphs:
            if par.text.startswith("№"):
                par.text = ''
                self.add_run(par, "№ карти___________________             Прізвище, ім’я, по батькові хворого  ")
                fio_run = self.add_run(par, self.patient.fio)
                fio_run.underline = True
                fio_run.bold = True
                self.add_run(par, " " * (65 - len(self.patient.fio)) + "№ Палати ")
                ward_run = self.add_run(par, self.patient.ward)
                ward_run.underline = True
                ward_run.bold = True
                break

    def set_mode_and_diet(self):
        table = self.document.tables[0]
        #mode_cell_info = self.get_cell_by_value_from_table(table, "Режим")
        mode_cell = table.cell(5, 1)  # Режим
        diet_cell = table.cell(6, 1)  # Дієта
        mode_cell.text = self.patient.mode
        diet_cell.text = self.patient.diet

    def set_date_range(self):
        max_duration = max([prescription.duration for prescription in self.patient.prescriptions])
        today = datetime.datetime.today()
        row = self.document.tables[0].rows[4]  # дата
        self.max_date_range = []
        for i in range(max_duration):
            index = i + 3  # дата.col_num + 1
            if index < len(row.cells):
                date = today + datetime.timedelta(days=i)
                self.max_date_range.append(date)
                cell = row.cells[index]
                self.set_text_to_cell(cell, date.strftime(day_month_date_format),  7)

    def set_prescriptions(self):
        table = self.document.tables[0]
        prescriptions = self.patient.prescriptions
        for i in range(len(prescriptions)):
            row_index = 1 + i * 2
            if row_index < len(table.rows):
                row = table.rows[row_index]
                prescription_cell = row.cells[0]
                presc = prescriptions[i]
                self.set_text_to_cell(prescription_cell, f"{presc.drug_with_dosage}  {presc.injection_method}  {presc.multiplicity}", 9)
                for j in range(presc.duration):
                    cell_index = 3 + j  # '+' start
                    if cell_index < len(row.cells):
                        plus_cell = row.cells[cell_index]
                        plus_cell.text = '  +'

    def set_surveys(self):
        consultations = 1
        table = self.document.tables[1]
        for survey in self.patient.surveys:
            if survey.row_num is not None and survey.col_num is not None:
                date_cell = table.cell(survey.row_num, survey.col_num+1)
                self.set_text_to_cell(date_cell, survey.date, 10, WD_TABLE_ALIGNMENT.CENTER)
            else:
                index = 36 + consultations  # Консультації row_num
                if index < len(table.rows):
                    needed_consult_row = table.rows[index]
                    consult_cell = needed_consult_row.cells[6]  # survey col_num
                    date_cell = needed_consult_row.cells[7]  # date col_num
                    self.set_text_to_cell(consult_cell, survey.name, 10)
                    self.set_text_to_cell(date_cell, survey.date, 10, WD_TABLE_ALIGNMENT.CENTER)
                    consultations += 1

    def add_nurse_tables(self):
        prescriptions = self.patient.prescriptions
        for date_index in range(len(self.max_date_range)):
            table = self.document.tables[-1]
            new_tbl = deepcopy(table._tbl)
            self.set_text_to_cell(table.cell(0, 0), self.patient.fio, 14, WD_TABLE_ALIGNMENT.CENTER)
            self.set_text_to_cell(table.cell(0, 2), f"№ {self.patient.ward}", 14, WD_TABLE_ALIGNMENT.CENTER)
            date_str = self.max_date_range[date_index].strftime(day_month_date_format)
            self.set_text_to_cell(table.cell(0, 3), date_str, 14, WD_TABLE_ALIGNMENT.CENTER)
            for i in range(len(prescriptions)):
                if date_index < prescriptions[i].duration:
                    injection_method = prescriptions[i].injection_method \
                                        if prescriptions[i].injection_method not in ['в/в стр.', 'в/в крап.'] \
                                        else ''
                    times = self.multiplicity_times_dict[prescriptions[i].multiplicity].split(',')
                    for time in times:
                        cell = self.get_cell_for_prescription(table, prescriptions[i], time)
                        if cell.text != '':
                            cell.add_paragraph().text = f"{prescriptions[i].drug_with_dosage} {injection_method}".strip()
                        else:
                            cell.text = f"{prescriptions[i].drug_with_dosage} {injection_method}".strip()
            if date_index != len(self.max_date_range) - 1:
                self.document.add_paragraph()._p.addnext(new_tbl)

    def get_cell_for_prescription(self, table, prescription, time) -> _Cell:
        row_index = self.get_cell_by_value_from_table(table, time).row_num
        column_index = self.get_cell_by_value_from_table(table, prescription.injection_method).col_num
        return table.cell(row_index, column_index)

    def save_document(self):
        name = f"{self.patient.fio.replace(' ', '_')}_{self.patient.ward}"
        count = 0
        for file in os.listdir("output_files/docx"):
            if name in file:
                count += 1
        if count > 0:
            self.document.save(f"output_files/docx/{name}_v{count + 1}.docx")
        else:
            self.document.save(f"output_files/docx/{name}.docx")

    @staticmethod
    def get_cell_by_value_from_table(table, value: str) -> CellInfo:
        for row_index in range(len(table.rows)):
            row = table.rows[row_index]
            for cell_index in range(len(row.cells)):
                cell = row.cells[cell_index]
                if value in cell.text:
                    return CellInfo(cell=cell, row_num=row_index, col_num=cell_index)

    @staticmethod
    def set_text_to_cell(cell, text, font_size, alignment=WD_TABLE_ALIGNMENT.LEFT, font_name='Times New Roman'):
        cell.text = text
        for par in cell.paragraphs:
            par.alignment = alignment
            for run in par.runs:
                run.font.size = Pt(font_size)
                run.font.name = font_name

    @staticmethod
    def add_run(par, text, font_name='Times New Roman', font_size=10):
        run = par.add_run(text)
        run.font.size = Pt(font_size)
        run.font.name = font_name
        return run

